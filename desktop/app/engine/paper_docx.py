"""Paper → docx: source/utils/markdownParser.js の cleanMarkdown + generateDocxBlob と同じ方針。

- Markdown を **HTML 経由にしない**。行単位で docx を組み立てる。
- タスク行は **Unicode ☐/☑** を先頭に付与（Google ドキュメントが Word の w14:checkbox をネイティブチェックリストに変換しないことが多いため）。
"""
from __future__ import annotations

import io
import logging
import re
from io import BytesIO

import httpx
from ..dropbox_oauth_refresh import dropbox_request_with_token_refresh
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

# markdownParser.js の parseInline と同等（Python re は \3 後方参照が並びによって無効になるため (?P=bd) を使用）
_INLINE_RE = re.compile(
    r"\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\)"
    r"|(?P<bd>\*\*|__)(?!\s)(?P<btxt>.*?)(?!\s)(?P=bd)"
    r"|(?:\b|_)\*(?!\s)(?P<ita>.*?)(?!\s)\*(?:\b|_)"
    r"|(?:\b|_)_(?!\s)(?P<itb>.*?)(?!\s)_(?:\b|_)",
    re.DOTALL,
)

# テーブル区切り行（GFM）
_TABLE_SEP_RE = re.compile(
    r"^\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)*\|?$"
)

# チェックリスト行: - [ ] / [x] 等。括弧内は空=未チェック、x/X=チェック済み（半角スペース複数も可）
_CHECK_LINE_RE = re.compile(r"^\s*(?:[-*+]\s+|\d+\.\s+)?\[\s*([xX]?)\s*\]\s*(.*)$")


def _paragraph_with_checkbox_and_label(
    paragraph: Paragraph, checked: bool, label: str
) -> None:
    """Google ドキュメントは Word の w14:checkbox をネイティブチェックリストにしばしば変換しないため、☐/☑ を使う。"""
    paragraph._p.clear_content()
    paragraph.add_run("\u2611 " if checked else "\u2610 ")
    if label.strip():
        _add_inline_to_paragraph(paragraph, label)


def clean_markdown(markdown: str) -> str:
    """markdownParser.js の cleanMarkdown に相当。"""
    if not markdown:
        return ""
    processed = re.sub(r"\*{4,}", "", markdown)
    processed = re.sub(r"\*\*(\s*)\*\*", r"\1", processed)
    lines = processed.split("\n")
    result_lines: list[str] = []
    current_list_indent = -1
    for line in lines:
        trimmed = line.strip()
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s", line)
        is_image = re.match(r"^!\[.*\]\(.*\)$", trimmed)
        if list_match:
            current_list_indent = len(list_match.group(1))
            result_lines.append(line)
        elif is_image and current_list_indent >= 0:
            result_lines.append(" " * (current_list_indent + 4) + trimmed)
        elif trimmed != "" and not line.startswith(" ") and not line.startswith("\t"):
            current_list_indent = -1
            result_lines.append(line)
        else:
            result_lines.append(line)
    processed = "\n".join(result_lines)
    processed = re.sub(r"([^\n|])\n\|", r"\1\n\n|", processed)

    # 4. Dropbox Paper のチェックリスト（1行が | ... | で囲まれている場合）
    def _pipe_checklist(m: re.Match[str]) -> str:
        content = m.group(1)
        if re.search(r"\[[ xX]\]", content):
            multi_line = re.sub(
                r"<br>\s*(?:-\s*)?(\[[ xX]\])",
                r"\n- \1",
                content,
            )
            multi_line = multi_line.replace("<br>", "\n")
            return multi_line
        return m.group(0)

    processed = re.sub(
        r"^\|\s*([^|]+)\s*\|$",
        _pipe_checklist,
        processed,
        flags=re.MULTILINE,
    )
    return processed


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    """ExternalHyperlink に相当。"""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A73E8")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_inline_to_paragraph(paragraph: Paragraph, text: str) -> None:
    """parseInline に相当。"""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        gd = m.groupdict()
        if gd.get("link_text") is not None:
            _add_hyperlink(paragraph, gd["link_text"], gd["link_url"])
        elif gd.get("bd") is not None:
            run = paragraph.add_run(gd["btxt"])
            run.bold = True
        elif gd.get("ita") is not None:
            run = paragraph.add_run(gd["ita"])
            run.italic = True
        elif gd.get("itb") is not None:
            run = paragraph.add_run(gd["itb"])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _set_cell_borders(cell, color: str = "DFE1E5", sz: str = "8") -> None:
    """セル四辺の罫線（w:sz は 1/8 pt。8 ≒ 1pt）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    else:
        tc_borders.clear()
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), sz)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)
        tc_borders.append(edge_el)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """ヘッダー行などの背景色（HTML の #E8F0FE 相当は E8F0FE）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _add_thematic_break(doc: Document) -> None:
    """thematicBreak の簡易代替（下線のみ）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _flush_table_docx(
    doc: Document,
    table_rows_buffer: list[str],
) -> None:
    if len(table_rows_buffer) < 2:
        for row in table_rows_buffer:
            p = doc.add_paragraph()
            _add_inline_to_paragraph(p, row)
        return
    parsed_rows: list[list[str]] = []
    for row in table_rows_buffer:
        t = row.strip()
        if _TABLE_SEP_RE.match(t):
            continue
        cells = t.strip().lstrip("|").rstrip("|").split("|")
        parsed_rows.append([c.strip() for c in cells])
    if not parsed_rows:
        return
    ncols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, cells in enumerate(parsed_rows):
        row_cells = list(cells[:ncols])
        while len(row_cells) < ncols:
            row_cells.append("")
        for ci, cell_text in enumerate(row_cells):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_to_paragraph(p, cell_text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_shading(cell, "E8F0FE")


async def markdown_to_docx_bytes(
    client: httpx.AsyncClient,
    token_ref: list[str],
    refresh: str | None,
    title: str,
    markdown_text: str,
) -> bytes:
    """generateDocxBlob と同じアルゴリズムでバイト列を返す。"""
    cleaned = clean_markdown(markdown_text)
    lines = cleaned.split("\n")
    doc = Document()
    if title:
        doc.core_properties.title = title

    table_rows_buffer: list[str] = []
    in_table = False

    async def fetch_image(url: str) -> bytes | None:
        if not url.startswith(("https://", "http://")):
            return None
        try:
            if (
                "paper-attachments.dropbox.com" in url
                or "paper-attachments.dropboxusercontent.com" in url
            ):
                r = await dropbox_request_with_token_refresh(
                    client,
                    token_ref,
                    refresh,
                    lambda tok: client.get(
                        url,
                        headers={
                            "User-Agent": "Mozilla/5.0",
                            "Authorization": f"Bearer {tok}",
                        },
                        timeout=60.0,
                        follow_redirects=True,
                    ),
                )
            else:
                r = await client.get(
                    url,
                    headers={"User-Agent": "Mozilla/5.0"},
                    timeout=60.0,
                    follow_redirects=True,
                )
        except Exception as e:
            logger.warning("paper_docx image fetch error url=%s: %s", url, e)
            return None
        if r.status_code != 200:
            return None
        return r.content

    for line in lines:
        is_table_line = line.strip().startswith("|")
        if is_table_line:
            in_table = True
            table_rows_buffer.append(line)
            continue
        if in_table:
            _flush_table_docx(doc, table_rows_buffer)
            table_rows_buffer = []
            in_table = False

        trimmed = line.strip()
        if not trimmed:
            doc.add_paragraph("")
            continue

        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", trimmed)
        if img_m:
            url = img_m.group(2)
            data = await fetch_image(url)
            if data:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(10)
                    run = p.add_run()
                    run.add_picture(BytesIO(data), width=Inches(5), height=Inches(3))
                except Exception:
                    doc.add_paragraph(f"[画像を docx に埋め込めません: {url}]")
            else:
                doc.add_paragraph(f"[画像取得失敗: {url}]")
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)
            _add_inline_to_paragraph(p, line[2:].strip())
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            p.paragraph_format.space_before = Pt(15)
            p.paragraph_format.space_after = Pt(8)
            _add_inline_to_paragraph(p, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 3"]
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(5)
            _add_inline_to_paragraph(p, line[4:].strip())
        elif re.match(r"^-{3,}$", trimmed):
            _add_thematic_break(doc)
        elif _CHECK_LINE_RE.match(line):
            cm = _CHECK_LINE_RE.match(line)
            assert cm is not None
            checked = (cm.group(1) or "").lower() == "x"
            label = cm.group(2)
            indent_match = re.match(r"^(\s*)", line)
            indent_level = (
                len(indent_match.group(1)) // 4 if indent_match else 0
            )
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if indent_level > 0:
                p.paragraph_format.left_indent = Pt(indent_level * 36)
            _paragraph_with_checkbox_and_label(p, checked, label)
        elif re.match(r"^\s*[-*] ", line):
            indent_match = re.match(r"^(\s*)", line)
            indent_level = (
                len(indent_match.group(1)) // 4 if indent_match else 0
            )
            body = re.sub(r"^\s*[-*] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            if indent_level > 0:
                p.paragraph_format.left_indent = Pt(indent_level * 36)
            _add_inline_to_paragraph(p, body)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_inline_to_paragraph(p, line)

    if in_table:
        _flush_table_docx(doc, table_rows_buffer)

    # Normal スタイルの東アジアフォント（既存の見た目に寄せる）
    try:
        style = doc.styles["Normal"]
        style.font.name = "Hiragino Kaku Gothic ProN"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Kaku Gothic ProN")
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
